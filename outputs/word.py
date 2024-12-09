from docx import Document
from docx.shared import Inches, Pt

# Initialize a Word document
doc = Document()

# Path to your image folder
image_folder = "."

# Loop to add images with text
for i in range(1, 51):  # Adjust range based on your images
    paragraph = doc.add_paragraph()
    
    # Add text with font size 24
    run = paragraph.add_run(f"Q{i} output: ")
    run.font.size = Pt(24)
    
    # Add image
    paragraph.add_run().add_picture(f"{image_folder}/{i}.png", width=Inches(3))  # Adjust width as needed

    # Add spacing between entries
    doc.add_paragraph("")  # Empty paragraph for spacing

# Save the document
doc.save("output.docx")
